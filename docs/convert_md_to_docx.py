#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Markdown to Word转换脚本
将SOP Markdown文档转换为Word文档,并插入截图
"""

import re
import os
from pathlib import Path

print("正在检查依赖库...")

# 检查并安装必要的库
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    print("[OK] python-docx installed")
except ImportError:
    print("[Installing] python-docx...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    print("[OK] python-docx installed")

try:
    import markdown
    print("[OK] markdown installed")
except ImportError:
    print("[Installing] markdown...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'markdown'])
    import markdown
    print("[OK] markdown installed")

def parse_markdown_to_word(md_file, output_file):
    """
    将Markdown文件转换为Word文档
    """
    print(f"\n开始转换: {md_file}")

    # 创建Word文档
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # 按行处理
    lines = content.split('\n')

    in_code_block = False
    code_block_content = []

    for line in lines:
        # 处理代码块
        if line.startswith('```'):
            if in_code_block:
                # 代码块结束
                if code_block_content:
                    code_text = '\n'.join(code_block_content)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Intense Quote'
                code_block_content = []
                in_code_block = False
            else:
                # 代码块开始
                in_code_block = True
            continue

        if in_code_block:
            code_block_content.append(line)
            continue

        # 处理标题
        if line.startswith('# '):
            text = line.lstrip('# ').strip()
            heading = doc.add_heading(text, level=1)
        elif line.startswith('## '):
            text = line.lstrip('# ').strip()
            heading = doc.add_heading(text, level=2)
        elif line.startswith('### '):
            text = line.lstrip('# ').strip()
            heading = doc.add_heading(text, level=3)
        elif line.startswith('#### '):
            text = line.lstrip('# ').strip()
            heading = doc.add_heading(text, level=4)

        # 处理图片引用
        elif '![' in line and '](' in line:
            # 提取图片路径
            match = re.search(r'!\[.*?\]\((.*?)\)', line)
            if match:
                img_path = match.group(1)
                # 转换相对路径
                if img_path.startswith('./'):
                    img_path = img_path.replace('./', '')

                full_img_path = os.path.join(os.path.dirname(md_file), '..', img_path)

                if os.path.exists(full_img_path):
                    try:
                        doc.add_picture(full_img_path, width=Inches(5.5))
                        p = doc.add_paragraph(f'图片: {os.path.basename(img_path)}')
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        print(f"  [OK] Image inserted: {img_path}")
                    except Exception as e:
                        print(f"  [ERROR] Cannot insert image {img_path}: {e}")
                        doc.add_paragraph(f'[Image path: {img_path}]')
                else:
                    print(f"  [WARNING] Image not found: {full_img_path}")
                    doc.add_paragraph(f'[Image path: {img_path}]')

        # 处理表格 (简化处理)
        elif line.strip().startswith('|') and '|' in line:
            # 简单的表格处理
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                # 检查是否是分隔行
                if all(cell.replace('-', '').strip() == '' for cell in cells):
                    continue
                # 这里简化处理,将表格内容作为段落
                doc.add_paragraph(' | '.join(cells))

        # 处理列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip().lstrip('-* ').strip()
            doc.add_paragraph(text, style='List Bullet')

        elif re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            doc.add_paragraph(text, style='List Number')

        # 处理普通段落
        elif line.strip():
            # 移除Markdown格式
            text = line
            # 移除粗体标记
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # 移除代码标记
            text = re.sub(r'`(.*?)`', r'\1', text)
            # 移除链接标记
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)

            if text.strip():
                doc.add_paragraph(text)

    # 保存文档
    doc.save(output_file)
    print(f"\n[OK] Conversion completed!")
    print(f"  Output file: {output_file}")
    print(f"  File size: {os.path.getsize(output_file) / 1024:.2f} KB")

if __name__ == '__main__':
    # 脚本所在目录
    script_dir = Path(__file__).parent

    # 输入文件
    md_file = script_dir / '非AI类工具网站开发SOP.md'

    # 输出文件
    output_file = script_dir / '非AI类工具网站开发SOP.docx'

    if not md_file.exists():
        print(f"[ERROR] Input file not found: {md_file}")
        exit(1)

    # 转换
    parse_markdown_to_word(str(md_file), str(output_file))

    print(f"\n[DONE] Word document saved to:")
    print(f"  {output_file.absolute()}")
